"""Synthesizer: build structured notes using LLM and templates.

Three modes:
1. API LLM MODE - Uses free API LLMs (Groq/Ollama/Together) - HIGH QUALITY
2. LOCAL LLM MODE - Uses Hugging Face models (Flan-T5, etc.) - LOW QUALITY
3. TEMPLATE MODE (default) - Fast, structured, instant output - RELIABLE

Includes historical context from previous periods.
"""
from __future__ import annotations
from typing import List, Dict, Optional
import re
from datetime import datetime, timedelta
import sqlite3
import os

from watcher.agents.llm_adapter import LocalLLMAdapter
try:
    from watcher.agents.llm_api_adapter import APILLMAdapter, get_recommended_adapter
except ImportError:
    APILLMAdapter = None
    get_recommended_adapter = None

try:
    from watcher.analysis.history import HistoricalAnalyzer
except ImportError:
    HistoricalAnalyzer = None


class Synthesizer:
    def __init__(self, 
                 model_name: str = None, 
                 adapter = None, 
                 use_llm: bool = False,
                 use_api_llm: bool = False,
                 api_provider: str = 'groq',
                 api_model: str = 'openai/gpt-oss-120b',
                 api_key: str = None):
        """
        Args:
            model_name: HF model to use (only if use_llm=True). If None, loads from config.
            adapter: Optional adapter instance (LocalLLMAdapter or APILLMAdapter)
            use_llm: If True, use local HF model (Flan-T5, etc.) - NOT RECOMMENDED
            use_api_llm: If True, use API-based LLM (Groq/Gemini/Ollama) - RECOMMENDED for quality
            api_provider: 'groq' (default), 'gemini', 'ollama', or 'together'
            api_model: Specific model to use (default: GPT-OSS 120B via Groq)
            api_key: API key for provider (or set GROQ_API_KEY/GEMINI_API_KEY env var)
        
        Priority:
        1. If use_api_llm=True -> Use Groq/Gemini/Ollama (HIGH QUALITY, FREE)
        2. If use_llm=True -> Use local Flan-T5 (LOW QUALITY, but works offline)
        3. Otherwise -> Use template mode (RELIABLE, INSTANT)
        
        Examples:
            # Best size (free GPT-OSS 120B via Groq) - DEFAULT
            synth = Synthesizer(use_api_llm=True)
            
            # Best quality (Gemini 2.0 Flash, 1M context)
            synth = Synthesizer(use_api_llm=True, api_provider='gemini', api_model='gemini-2.0-flash-exp')
            
            # Fast alternative (Llama 3.3 70B)
            synth = Synthesizer(use_api_llm=True, api_model='llama-3.3-70b-versatile')
            
            # Local private (requires Ollama installed)
            synth = Synthesizer(use_api_llm=True, api_provider='ollama')
            
            # Fast template (default, recommended if no API)
            synth = Synthesizer()
        """
        self.use_api_llm = use_api_llm
        self.use_llm = use_llm
        self.adapter = adapter
        
        # Mode 1: API LLM (best quality)
        if use_api_llm and APILLMAdapter:
            if adapter is None:
                try:
                    # Initialize adapter
                    self.adapter = APILLMAdapter(
                        provider=api_provider,
                        api_key=api_key,
                        model=api_model
                    )
                    
                    # Display appropriate message
                    if api_provider == 'groq':
                        if 'gpt-oss-120b' in api_model:
                            print(f"✅ Using Groq API (GPT-OSS 120B) for high-quality synthesis")
                        else:
                            print(f"✅ Using Groq API ({api_model}) for high-quality synthesis")
                    elif api_provider == 'gemini':
                        print(f"✅ Using Gemini API ({api_model}) for highest-quality synthesis")
                    else:
                        print(f"✅ Using {api_provider.upper()} API for synthesis")
                except Exception as e:
                    print(f"⚠️ API LLM setup failed: {e}")
                    print("   Falling back to template mode")
                    self.use_api_llm = False
        
        # Mode 2: Local LLM (Flan-T5, etc.)
        elif use_llm:
            if adapter is None:
                try:
                    from watcher.config import load_config
                    config = load_config()
                    model_name = model_name or config.get('synthesizer_model', 'google/flan-t5-large')
                    task = config.get('synthesizer_task', 'text2text-generation')
                except:
                    model_name = model_name or 'google/flan-t5-large'
                    task = 'text2text-generation'
                
                self.adapter = LocalLLMAdapter(model_name=model_name, task=task)
                print(f"⚠️ Using local {model_name} (may produce low-quality output)")
            
        # Mode 3: Template (default, reliable)
        else:
            print("✅ Using template mode for reliable structured output")
        
        self.model_name = model_name or 'template'

    def synthesize(self, topic: str, period: str, context: str, items: List[Dict], max_new_tokens: int = 6000, db_path: str = "watcher.db") -> str:
        """Generate structured note - fast template mode by default, optional LLM.
        
        Args:
            topic: Subject of analysis
            period: Time period label
            context: Background context
            items: List of items to analyze
            max_new_tokens: Max tokens for LLM output
            db_path: Path to database for historical context
        """
        if not items:
            return self._generate_empty_note(topic, period)
        
        # Get historical context if available
        historical_context = ""
        if HistoricalAnalyzer:
            try:
                analyzer = HistoricalAnalyzer(db_path)
                prev_summary = analyzer.get_previous_period_summary(current_period_days=7)
                if prev_summary['item_count'] > 0:
                    historical_context = f"\n## Contexte de la période précédente\n{prev_summary['summary']}"
            except:
                pass  # Silently fail if historical analysis unavailable
        
        # Prioritize API LLM > Local LLM > Template
        if self.use_api_llm and self.adapter:
            return self._synthesize_with_llm(topic, period, context, items, max_new_tokens, historical_context)
        elif self.use_llm and self.adapter:
            return self._synthesize_with_llm(topic, period, context, items, max_new_tokens, historical_context)
        else:
            return self._synthesize_with_template(topic, period, context, items, historical_context)
    
    def _synthesize_with_template(self, topic: str, period: str, context: str, items: List[Dict], historical_context: str = "") -> str:
        """Fast template-based synthesis (instant, always structured)."""
        
        # Extract key information from items
        titles = [item.get('title', 'Sans titre') for item in items]
        sources = [(item.get('source', 'Unknown'), item.get('url', '#')) for item in items]
        summaries = [item.get('summary', item.get('content', '')[:200]) for item in items]
        
        # Build structured note
        lines = []
        lines.append("# RÉSUMÉ EXÉCUTIF")
        lines.append(self._build_summary(titles, len(items)))
        
        lines.append("\n# CONTEXTE")
        lines.append(f"Sujet: {topic}")
        lines.append(f"Période: {period}")
        lines.append(f"Nombre d'articles analysés: {len(items)}")
        lines.append(f"Description: {context}")
        if historical_context:
            lines.append(historical_context)
        
        lines.append("\n# NOUVEAUTÉS PRINCIPALES")
        for i, title in enumerate(titles[:5], 1):
            lines.append(f"• {title} [{i}]")
        
        lines.append("\n# ANALYSE ET IMPLICATIONS")
        lines.append(self._build_analysis(items))
        
        lines.append("\n# SOURCES")
        for i, (item, (source, url)) in enumerate(zip(items, sources), 1):
            title = item.get('title', 'Sans titre')[:80]
            date = item.get('published', '')
            lines.append(f"[{i}] {title}")
            lines.append(f"    Source: {source}")
            if date:
                lines.append(f"    Date: {date}")
            if url and url != '#':
                lines.append(f"    URL: {url}")
        
        return "\n".join(lines)
    
    def _build_summary(self, titles: List[str], count: int) -> str:
        """Generate intelligent executive summary from titles."""
        if not titles:
            return "Aucun article analysé."
        
        # Extract common themes from titles
        themes = self._extract_themes(titles)
        
        if themes:
            theme_str = ", ".join(themes[:3])
            summary = f"Cette période a révélé {count} développements significatifs, "
            summary += f"principalement autour de: {theme_str}. "
        else:
            summary = f"Analyse de {count} articles récents. "
        
        # Highlight top story
        if titles:
            summary += f"Point majeur: {titles[0][:100]}{'...' if len(titles[0]) > 100 else ''}."
        
        return summary
    
    def _build_analysis(self, items: List[Dict]) -> str:
        """Generate intelligent analysis and implications from items."""
        lines = []
        
        # Count unique sources
        sources = set(item.get('source', 'Unknown') for item in items)
        lines.append(f"**Couverture**: {len(items)} articles provenant de {len(sources)} source(s) distincte(s).")
        
        # Categorize items if they have categories
        categories = {}
        for item in items:
            cat = item.get('category', 'Non classé')
            categories[cat] = categories.get(cat, 0) + 1
        
        if len(categories) > 1:
            lines.append("\n**Répartition par type**:")
            for cat, count in sorted(categories.items(), key=lambda x: x[1], reverse=True):
                lines.append(f"  • {cat.title()}: {count} article(s)")
        
        # Identify patterns/trends
        keywords = {}
        for item in items:
            text = (item.get('title', '') + ' ' + item.get('summary', '')).lower()
            tech_keywords = ['ai', 'artificial intelligence', 'machine learning', 'deep learning', 
                           'neural network', 'chatgpt', 'llm', 'generative', 'robotics',
                           'cloud', 'security', 'blockchain', 'quantum']
            for keyword in tech_keywords:
                if keyword in text:
                    keywords[keyword] = keywords.get(keyword, 0) + 1
        
        if keywords:
            top_keywords = sorted(keywords.items(), key=lambda x: x[1], reverse=True)[:4]
            lines.append("\n**Tendances technologiques détectées**:")
            for keyword, count in top_keywords:
                pct = int(100 * count / len(items))
                lines.append(f"  • {keyword.title()}: {count} mentions ({pct}% des articles)")
        
        # Add implications
        lines.append("\n**Implications**:")
        if len(items) > 5:
            lines.append(f"Volume élevé d'activité ({len(items)} articles) suggère une période de développements importants.")
        else:
            lines.append("Période de veille standard avec quelques développements notables.")
        
        if keywords:
            dominant = max(keywords.items(), key=lambda x: x[1])[0]
            lines.append(f"Le thème '{dominant}' domine l'actualité et mérite une attention particulière.")
        
        lines.append("La surveillance continue est recommandée pour suivre l'évolution de ces tendances.")
        
        return "\n".join(lines)

    def _synthesize_with_llm(self, topic: str, period: str, context: str, items: List[Dict], max_new_tokens: int, historical_context: str = "") -> str:
        """LLM-based synthesis with minimal validation.
        
        NO TEMPLATE FALLBACK - we trust the LLM output.
        Only falls back on actual API/generation errors.
        """
        # Cap items sent to LLM for speed — items are already ranked by relevance
        try:
            from watcher.config import load_config
            _cfg = load_config() or {}
            max_synthesis_items = int(_cfg.get("max_synthesis_items", 25))
        except Exception:
            max_synthesis_items = 25
        if len(items) > max_synthesis_items:
            print(f"ℹ️  Capping LLM input: {len(items)} → {max_synthesis_items} items (set max_synthesis_items in config.yaml to change)")
            items = items[:max_synthesis_items]

        prompt = self._build_llm_prompt(topic, period, context, items, historical_context)
        
        try:
            raw_output = self.adapter.generate(
                prompt,
                max_new_tokens=max_new_tokens,
                do_sample=True,
                temperature=0.7,
                top_p=0.9
            )
        except TypeError:
            try:
                raw_output = self.adapter.generate(prompt, max_new_tokens=max_new_tokens)
            except Exception as e:
                # Only fallback on actual generation errors
                print(f"❌ LLM API error: {e}. Using template as emergency fallback.")
                return self._synthesize_with_template(topic, period, context, items, historical_context)
        except Exception as e:
            # Only fallback on actual generation errors
            print(f"❌ LLM API error: {e}. Using template as emergency fallback.")
            return self._synthesize_with_template(topic, period, context, items, historical_context)
        
        # Strip reasoning blocks from DeepSeek R1 and similar models
        raw_output = self._strip_thinking_blocks(raw_output)
        
        # Only reject if truly garbage (extremely lenient)
        if self._is_garbage_output(raw_output, prompt):
            print("⚠️ LLM produced empty output. Using template as emergency fallback.")
            return self._synthesize_with_template(topic, period, context, items, historical_context)
        
        # Light structure enforcement (only adds sources if missing)
        return self._enforce_structure(raw_output, items, topic, period, context, historical_context)
    
    def _build_llm_prompt(self, topic: str, period: str, context: str, items: List[Dict], historical_context: str = "") -> str:
        """Build prompt for LLM synthesis - concise, professional format."""
        items_text = self._format_items_for_llm(items)

        historical_note = f"\n\nPériode précédente : {historical_context}" if historical_context else ""

        return f"""Vous êtes un analyste en veille technologique. Rédigez une note de veille concise et professionnelle.

FORMAT OBLIGATOIRE — respectez exactement cette structure :

# VEILLE TECHNOLOGIQUE : {topic}
Période : {period}

## RÉSUMÉ EXÉCUTIF
2 à 3 phrases. Citez les 2-3 annonces majeures avec leurs noms précis et chiffres clés si disponibles. Terminez par la tendance dominante.

## RAPPEL DU CONTEXTE
2 à 3 phrases maximum. Que se passait-il la période précédente ? Comment le contexte a évolué ?{historical_note}

## NOUVEAUTÉS DE LA PÉRIODE

### 1. [Titre de l'article]
Un paragraphe de 3 à 5 phrases : ce qui a été annoncé, pourquoi c'est important, données chiffrées si disponibles.
Source : [Nom de la source] - [Date]

### 2. [Titre]
[Même format]
Source : [Nom] - [Date]

### 3. [Titre]
[Même format]
Source : [Nom] - [Date]

(Continuez pour chaque article pertinent)

## ANALYSE ET IMPLICATIONS
Un seul paragraphe de 4 à 6 phrases. Quels patterns émergent ? Quelles implications techniques et business ? Quelle direction pour l'industrie ?

## SOURCES ET RÉFÉRENCES
• [Source] - "Titre" - [Date]
• ...

---

RÈGLES :
- Chaque article : 3 à 5 phrases, pas de sous-sections, pas de listes internes
- Pas de répétition entre le résumé et les articles
- Chiffres et noms précis quand disponibles, sinon ne pas inventer
- Langue : français

DONNÉES :
Sujet : {topic}
Période : {period}
Contexte : {context}

Articles ({len(items)}) :
{items_text}

Rédigez la note maintenant :
"""
    
    def _format_items_for_llm(self, items: List[Dict]) -> str:
        """Format items for LLM processing with full content - focus on quality over quantity."""
        lines = []
        # Limit to top 10 items for deeper analysis (quality over quantity)
        items_to_process = items[:10]
        
        for i, item in enumerate(items_to_process, 1):
            title = item.get('title', '(sans titre)')
            url = item.get('url', '')
            src = item.get('source', '')
            date = item.get('published', '')
            content = item.get('content', item.get('summary', ''))
            
            lines.append(f"\n[Article {i}]")
            lines.append(f"Titre: {title}")
            lines.append(f"Source: {src}")
            if date:
                lines.append(f"Date: {date}")
            if url:
                lines.append(f"URL: {url}")
            
            # Include content for better analysis (increased for quality focus)
            if content and len(content) > 100:
                lines.append(f"Contenu complet:")
                lines.append(content[:3000])  # Increased to 3000 chars for deeper analysis
                if len(content) > 3000:
                    lines.append("... (tronqué pour longueur)")
            elif content:
                lines.append(f"Résumé: {content}")
            else:
                lines.append("(Contenu non disponible - analysez le titre)")
            
            lines.append("")
        
        if len(items) > 10:
            lines.append(f"\n(Note: {len(items) - 10} articles supplémentaires disponibles mais non détaillés - focus qualité sur les {len(items_to_process)} plus pertinents)")
        
        return "\n".join(lines)
    
    def _strip_thinking_blocks(self, text: str) -> str:
        """Remove <think>...</think> blocks from DeepSeek R1 and similar reasoning models."""
        import re
        # Remove all content between <think> and </think> tags (case insensitive, multiline)
        cleaned = re.sub(r'<think>.*?</think>', '', text, flags=re.IGNORECASE | re.DOTALL)
        # Also handle incomplete think blocks at the start
        cleaned = re.sub(r'^<think>.*?(?=\n\n|\n#|$)', '', cleaned, flags=re.IGNORECASE | re.DOTALL)
        return cleaned.strip()
    
    def _is_garbage_output(self, output: str, prompt: str) -> bool:
        """Detect if LLM output is garbage (just repeating the prompt).
        
        VERY LENIENT - only rejects truly empty or nonsensical output.
        DeepSeek R1 and other reasoning models should pass through.
        """
        # DeepSeek R1 and reasoning models often have <think> blocks - strip them first for validation
        import re
        cleaned_output = re.sub(r'<think>.*?</think>', '', output, flags=re.IGNORECASE | re.DOTALL).strip()
        
        # If output after stripping think blocks is too short, it's garbage
        if len(cleaned_output) < 30:
            return True
        
        # Check if output is just the instructions being repeated (very strict check)
        if 'IMPORTANT: ANALYSEZ' in cleaned_output.upper() and 'DONNEZ des insights' in cleaned_output:
            # Likely just repeating prompt
            return True
        
        # Otherwise, trust the LLM output
        return False
    
    def _extract_themes(self, titles: List[str]) -> List[str]:
        """Extract common themes from article titles."""
        # Common tech themes
        theme_keywords = {
            'Intelligence Artificielle': ['ai', 'artificial intelligence', 'machine learning', 'deep learning', 'neural', 'llm', 'chatgpt', 'gpt'],
            'Sécurité': ['security', 'cybersecurity', 'breach', 'hack', 'vulnerability', 'ransomware'],
            'Cloud': ['cloud', 'aws', 'azure', 'gcp', 'kubernetes', 'docker'],
            'Blockchain': ['blockchain', 'crypto', 'bitcoin', 'ethereum', 'web3'],
            'Développement': ['programming', 'developer', 'code', 'software', 'framework', 'library'],
            'Innovation': ['innovation', 'breakthrough', 'new technology', 'advancement'],
            'Données': ['data', 'analytics', 'big data', 'database'],
            'Robotique': ['robot', 'robotics', 'automation', 'drone']
        }
        
        detected_themes = {}
        combined_text = ' '.join(titles).lower()
        
        for theme, keywords in theme_keywords.items():
            count = sum(1 for kw in keywords if kw in combined_text)
            if count > 0:
                detected_themes[theme] = count
        
        # Return top themes sorted by frequency
        sorted_themes = sorted(detected_themes.items(), key=lambda x: x[1], reverse=True)
        return [theme for theme, _ in sorted_themes[:4]]
    
    def _enforce_structure(self, raw_output: str, items: List[Dict], topic: str = "", period: str = "", context: str = "", historical_context: str = "") -> str:
        """Light structure enforcement - only adds missing SOURCES section.
        
        VERY LENIENT - accepts DeepSeek R1's natural output style.
        No template fallback - we trust the LLM.
        """
        # Only add sources if completely missing - don't enforce other sections
        if "# SOURCES" not in raw_output and "## SOURCES" not in raw_output:
            sources_section = "\n\n# SOURCES\n"
            for i, item in enumerate(items, 1):
                title = item.get('title', 'Sans titre')[:60]
                url = item.get('url', 'N/A')
                sources_section += f"[{i}] {title} - {url}\n"
            raw_output += sources_section
        
        return raw_output
    
    def _generate_empty_note(self, topic: str, period: str) -> str:
        """Generate note when no items available."""
        return f"""# RÉSUMÉ EXÉCUTIF
Aucun article nouveau pour la période {period}.

# CONTEXTE
Sujet: {topic}
Période: {period}

# NOUVEAUTÉS PRINCIPALES
Aucune nouveauté détectée.

# ANALYSE ET IMPLICATIONS
Le système de veille continue de surveiller les sources. Aucun développement majeur identifié.

# SOURCES
Aucune source à rapporter.
"""
    
    def export_to_pdf(self, note_text: str, output_path: str, topic: str = "", period: str = "") -> str:
        """
        Export synthesis note to PDF with professional formatting.
        
        Args:
            note_text: The synthesis text to export
            output_path: Path where PDF will be saved
            topic: Optional topic for header
            period: Optional period for header
            
        Returns:
            Path to generated PDF file
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            from reportlab.pdfgen import canvas
            from datetime import datetime
        except ImportError:
            raise ImportError("reportlab required for PDF export. Install with: pip install reportlab")
        
        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=A4,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#1a1a1a',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#2c3e50',
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            textColor='#333333',
            spaceAfter=12,
            leading=14,
            fontName='Helvetica'
        )
        
        # Add title
        if topic:
            elements.append(Paragraph(f"Note de Veille: {topic}", title_style))
        else:
            elements.append(Paragraph("Note de Veille", title_style))
        
        # Add period and date
        meta_text = f"<i>Période: {period or 'N/A'} | Généré le: {datetime.now().strftime('%d/%m/%Y à %H:%M')}</i>"
        elements.append(Paragraph(meta_text, body_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Process note text
        lines = note_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 0.1*inch))
                continue
            
            # Headings (lines starting with #)
            if line.startswith('# '):
                heading_text = line.replace('# ', '').strip()
                elements.append(Paragraph(heading_text, heading_style))
            # Bullet points
            elif line.startswith('• ') or line.startswith('- '):
                bullet_text = line.replace('• ', '').replace('- ', '')
                elements.append(Paragraph(f"• {bullet_text}", body_style))
            # Regular text
            else:
                # Replace markdown bold **text** with <b>text</b>
                line = line.replace('**', '<b>', 1).replace('**', '</b>', 1)
                elements.append(Paragraph(line, body_style))
        
        # Build PDF
        doc.build(elements)
        
        return output_path
    
    def export_to_docx(self, note_text: str, output_path: str, topic: str = "", period: str = "") -> str:
        """
        Export synthesis note to Word document with professional formatting.
        
        Args:
            note_text: The synthesis text to export
            output_path: Path where DOCX will be saved
            topic: Optional topic for header
            period: Optional period for header
            
        Returns:
            Path to generated DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
        except ImportError:
            raise ImportError("python-docx required for Word export. Install with: pip install python-docx")
        
        # Create document
        doc = Document()
        
        # Add title
        if topic:
            title = doc.add_heading(f"Note de Veille: {topic}", 0)
        else:
            title = doc.add_heading("Note de Veille", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Période: {period or 'N/A'} | Généré le: {datetime.now().strftime('%d/%m/%Y à %H:%M')}").italic = True
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Process note text
        lines = note_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            # Headings (lines starting with #)
            if line.startswith('# '):
                heading_text = line.replace('# ', '').strip()
                heading = doc.add_heading(heading_text, level=1)
                heading_format = heading.runs[0].font
                heading_format.color.rgb = RGBColor(44, 62, 80)
            # Bullet points
            elif line.startswith('• ') or line.startswith('- '):
                bullet_text = line.replace('• ', '').replace('- ', '')
                doc.add_paragraph(bullet_text, style='List Bullet')
            # Regular text
            else:
                para = doc.add_paragraph()
                # Handle bold text **text**
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        para.add_run(part)
                    else:
                        para.add_run(part).bold = True
        
        # Save document
        doc.save(output_path)
        
        return output_path
